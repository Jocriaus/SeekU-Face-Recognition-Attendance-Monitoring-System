import tkinter as Tk
import pandas as pd
import query_mod as qry
import os
import docx
import docx2pdf
import Treeview_table_mod as tbl
import tkinter.messagebox as messbx

class excelClass:
    def __init__(self, master=None):
        self.sql_query = qry.dbQueries()
        self.treeview = tbl.TreeviewGUI()

    def save_student(self, filepath, filename, date1, date2):
        data, columns = self.sql_query.sort_student_report_bydate_excel(date1, date2)
        columns = ['Student No.', 'First Name', 'Last Name','Program' ,'Section' ,'Attendance Date', 'Time in', 'Time out' ]
        if data and columns:
            df = pd.DataFrame(data=data, columns=columns)
            writer = pd.ExcelWriter(
                filepath + "/" + filename + ".xlsx", engine="xlsxwriter"
            )
            start_column = 2
            df.to_excel(
                writer,
                sheet_name="Sheet1",
                startrow=12,
                startcol=start_column,
                header=True,
                index=False,
                na_rep="NaN",
            )
            starting_date = str(date1)[0:10]
            end_date = str(date2)[0:10]
 
            worksheet = writer.sheets["Sheet1"]
            worksheet.write('E9', starting_date)
            worksheet.write('G9', end_date)  
            worksheet.insert_image("C1", ".\SeekU\STI College Balagtas Logo medium.png")
            worksheet.insert_image("E1", ".\SeekU\Student Report.png")
            worksheet.insert_image("E6", ".\SeekU\Dates.png")
            worksheet.insert_image("I1", ".\SeekU\SeekU small.png")
            for column in df:
                column_length = max(df[column].astype(str).map(len).max(), len(column))
                col_idx = df.columns.get_loc(column)
                col_idx = col_idx + start_column
                worksheet.set_column(col_idx, col_idx, column_length)

            writer.close()
        else:
            messbx.showwarning(
                "Warning", "There are no records of attendance for the selected date."
            )

    def save_personnel(self, filepath, filename, date1, date2):
        data, columns = self.sql_query.sort_personnel_report_bydate_excel(date1, date2)
        columns = ['Personnel No.', 'First Name', 'Last Name','Personnel Type' ,'Attendance Date', 'Time in', 'Time out' ]
        if data and columns:
            df = pd.DataFrame(data=data, columns=columns)
            writer = pd.ExcelWriter(
                filepath + "/" + filename + ".xlsx", engine="xlsxwriter"
            )
            start_column = 3
            df.to_excel(
                writer,
                sheet_name="Sheet1",
                startrow=12,
                startcol=start_column,
                header=True,
                index=False,
                na_rep="NaN",
            )
            starting_date = str(date1)[0:10]
            end_date = str(date2)[0:10]
 
            worksheet = writer.sheets["Sheet1"]
            worksheet.write('E9', starting_date)
            worksheet.write('G9', end_date)  
            worksheet.insert_image("C1", ".\SeekU\STI College Balagtas Logo medium.png")
            worksheet.insert_image("E1", ".\SeekU\Personnel Report.png")
            worksheet.insert_image("E6", ".\SeekU\Dates.png")
            worksheet.insert_image("J1", ".\SeekU\SeekU small.png")
            for column in df:
                column_length = max(df[column].astype(str).map(len).max(), len(column))
                col_idx = df.columns.get_loc(column)
                col_idx = col_idx + start_column
                worksheet.set_column(col_idx, col_idx, column_length)
            writer.close()
        else:
            messbx.showwarning(
                "Warning", "There are no records of attendance for the selected date."
            )


    def save_visitor(self, filepath, filename, date1, date2):
        data, columns = self.sql_query.sort_visitor_report_bydate_excel(date1, date2)
        columns = ['Visitor No.', 'First Name', 'Last Name','Attendance Date', 'Time in', 'Time out' ]
        if data and columns:
            df = pd.DataFrame(data=data, columns=columns)
            writer = pd.ExcelWriter(
                filepath + "/" + filename + ".xlsx", engine="xlsxwriter"
            )
            start_column = 3
            df.to_excel(
                writer,
                sheet_name="Sheet1",
                startrow=12,
                startcol=start_column,
                header=True,
                index=False,
                na_rep="NaN",
            )
            starting_date = str(date1)[0:10]
            end_date = str(date2)[0:10]
 
            worksheet = writer.sheets["Sheet1"]
            worksheet.write('E9', starting_date)
            worksheet.write('H9', end_date)  
            worksheet.insert_image("C1", ".\SeekU\STI College Balagtas Logo medium.png")
            worksheet.insert_image("E1", ".\SeekU\Visitor Report.png")
            worksheet.insert_image("E6", ".\SeekU\Dates.png")
            worksheet.insert_image("I1", ".\SeekU\SeekU small.png")
            for column in df:
                column_length = max(df[column].astype(str).map(len).max(), len(column))
                col_idx = df.columns.get_loc(column)
                col_idx = col_idx + start_column
                worksheet.set_column(col_idx, col_idx, column_length)
            writer.close()
        else:
            messbx.showwarning(
                "Warning", "There are no records of attendance for the selected date."
            )


    def print_student(self, filepath, filename, date1, date2):
        data, columns = self.sql_query.sort_student_report_bydate_excel(date1, date2)
        columns = ['Student No.', 'First Name', 'Last Name','Program' ,'Section' ,'Attendance Date', 'Time in', 'Time out' ]
        if data and columns:
            df = pd.DataFrame(data=data, columns=columns)
            writer = pd.ExcelWriter(
                filepath + "/" + filename + ".xlsx", engine="xlsxwriter"
            )
            start_column = 2
            df.to_excel(
                writer,
                sheet_name="Sheet1",
                startrow=12,
                startcol=start_column,
                header=True,
                index=False,
                na_rep="NaN",
            )
            starting_date = str(date1)[0:10]
            end_date = str(date2)[0:10]
 
            worksheet = writer.sheets["Sheet1"]
            worksheet.write('E9', starting_date)
            worksheet.write('G9', end_date)  
            worksheet.insert_image("C1", ".\SeekU\STI College Balagtas Logo medium.png")
            worksheet.insert_image("E1", ".\SeekU\Student Report.png")
            worksheet.insert_image("E6", ".\SeekU\Dates.png")
            worksheet.insert_image("I1", ".\SeekU\SeekU small.png")
            for column in df:
                column_length = max(df[column].astype(str).map(len).max(), len(column))
                col_idx = df.columns.get_loc(column)
                col_idx = col_idx + start_column
                worksheet.set_column(col_idx, col_idx, column_length)

            writer.close()
            os.startfile(filepath + "/" + filename + ".xlsx")
        else:
            messbx.showwarning(
                "Warning", "There are no records of attendance for the selected date."
            )

    def print_personnel(self, filepath, filename, date1, date2):
        data, columns = self.sql_query.sort_personnel_report_bydate_excel(date1, date2)
        columns = ['Personnel No.', 'First Name', 'Last Name','Personnel Type' ,'Attendance Date', 'Time in', 'Time out' ]
        if data and columns:
            df = pd.DataFrame(data=data, columns=columns)
            writer = pd.ExcelWriter(
                filepath + "/" + filename + ".xlsx", engine="xlsxwriter"
            )
            start_column = 3
            df.to_excel(
                writer,
                sheet_name="Sheet1",
                startrow=12,
                startcol=start_column,
                header=True,
                index=False,
                na_rep="NaN",
            )
            starting_date = str(date1)[0:10]
            end_date = str(date2)[0:10]
 
            worksheet = writer.sheets["Sheet1"]
            worksheet.write('F9', starting_date)
            worksheet.write('G9', end_date)  
            worksheet.insert_image("C1", ".\SeekU\STI College Balagtas Logo medium.png")
            worksheet.insert_image("F1", ".\SeekU\Personnel Report.png")
            worksheet.insert_image("F6", ".\SeekU\Dates.png")
            worksheet.insert_image("J1", ".\SeekU\SeekU small.png")
            for column in df:
                column_length = max(df[column].astype(str).map(len).max(), len(column))
                col_idx = df.columns.get_loc(column)
                col_idx = col_idx + start_column
                worksheet.set_column(col_idx, col_idx, column_length)
            os.startfile(filepath + "/" + filename + ".xlsx")
        else:
            messbx.showwarning(
                "Warning", "There are no records of attendance for the selected date."
            )

    def print_visitor(self, filepath, filename, date1, date2):
        data, columns = self.sql_query.sort_visitor_report_bydate_excel(date1, date2)
        columns = ['Visitor No.', 'First Name', 'Last Name','Attendance Date', 'Time in', 'Time out' ]
        if data and columns:
            df = pd.DataFrame(data=data, columns=columns)
            writer = pd.ExcelWriter(
                filepath + "/" + filename + ".xlsx", engine="xlsxwriter"
            )
            start_column = 3
            df.to_excel(
                writer,
                sheet_name="Sheet1",
                startrow=12,
                startcol=start_column,
                header=True,
                index=False,
                na_rep="NaN",
            )
            starting_date = str(date1)[0:10]
            end_date = str(date2)[0:10]
 
            worksheet = writer.sheets["Sheet1"]
            worksheet.write('E9', starting_date)
            worksheet.write('H9', end_date)  
            worksheet.insert_image("C1", ".\SeekU\STI College Balagtas Logo medium.png")
            worksheet.insert_image("E1", ".\SeekU\Visitor Report.png")
            worksheet.insert_image("E6", ".\SeekU\Dates.png")
            worksheet.insert_image("I1", ".\SeekU\SeekU small.png")
            for column in df:
                column_length = max(df[column].astype(str).map(len).max(), len(column))
                col_idx = df.columns.get_loc(column)
                col_idx = col_idx + start_column
                worksheet.set_column(col_idx, col_idx, column_length)
            os.startfile(filepath + "/" + filename + ".xlsx")
        else:
            messbx.showwarning(
                "Warning", "There are no records of attendance for the selected date."
            )


class docxClass:
    def __init__(self, master=None):
        self.sql_query = qry.dbQueries()
        self.treeview = tbl.TreeviewGUI()
        self.main = Tk.Toplevel()
        self.treeview.student_report_treeview(self.main)
        self.treeview.visitor_report_treeview(self.main)
        self.treeview.personnel_report_treeview(self.main)
        self.main.withdraw()

    # JOCRIAUS--------------------------------------------------------------------------
    def print_doc_student(self, filepath, filename, date1, date2):
        # populate the report tree
        cont_inue = self.treeview.populate_student_report_bydate(date1, date2)
        # open an existing Word document
        if cont_inue:
            doc = docx.Document(".\Documents\Student Report Template.docx")

            column_names = [
                "Student_No",
                "Firstname",
                "Lastname",
                "Program",
                "Section",
                "Date",
                "Time In",
                "Time Out",
            ]

            table = doc.add_table(rows=1, cols=len(column_names))
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells

            for i in range(len(column_names)):
                hdr_cells[i].text = column_names[i]

            treeview_data = []
            # get the table data of student
            for child in self.treeview.student_report_tree.get_children():
                values = self.treeview.student_report_tree.item(child)["values"]
                if values not in treeview_data:
                    treeview_data.append(values)

            # Inserts the table data of student
            for row in treeview_data:
                row_cells = table.add_row().cells
                for i in range(len(column_names)):
                    row_cells[i].text = str(row[i])

            # saves the doc to a new file path
            doc.save(filepath + "/" + filename + ".docx")
            os.startfile(filepath + "/" + filename + ".docx")

    def save_doc_student(self, filepath, filename, date1, date2):
        # populate the report tree
        cont_inue = self.treeview.populate_student_report_bydate(date1, date2)
        if cont_inue:
            # open an existing Word document
            doc = docx.Document(".\Documents\Student Report Template.docx")

            column_names = [
                "Student_No",
                "Firstname",
                "Lastname",
                "Program",
                "Section",
                "Date",
                "Time In",
                "Time Out",
            ]

            table = doc.add_table(rows=1, cols=len(column_names))
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells

            for i in range(len(column_names)):
                hdr_cells[i].text = column_names[i]

            treeview_data = []
            # get the table data of student
            for child in self.treeview.student_report_tree.get_children():
                values = self.treeview.student_report_tree.item(child)["values"]
                if values not in treeview_data:
                    treeview_data.append(values)

            # Inserts the table data of student
            for row in treeview_data:
                row_cells = table.add_row().cells
                for i in range(len(column_names)):
                    row_cells[i].text = str(row[i])

            # saves the doc to a new file path
            doc.save(filepath + "/" + filename + ".docx")

    # JOCRIAUS--------------------------------------------------------------------------

    def print_doc_personnel(self, filepath, filename, date1, date2):
        # populate the report tree
        cont_inue = self.treeview.populate_personnel_report_bydate(date1, date2)
        if cont_inue:
            # open an existing Word document
            doc = docx.Document(".\Documents\Personnel Report Template.docx")

            column_names = [
                "Personnel Number",
                "First Name",
                "Last Name",
                "Personnel Type",
                "Date",
                "Time In",
                "Time Out",
            ]
            table = doc.add_table(rows=1, cols=len(column_names))
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells

            for i in range(len(column_names)):
                hdr_cells[i].text = column_names[i]

            treeview_data = []
            # get the table data of student
            for child in self.treeview.personnel_report_tree.get_children():
                values = self.treeview.personnel_report_tree.item(child)["values"]
                if values not in treeview_data:
                    treeview_data.append(values)

            # Inserts the table data of student
            for row in treeview_data:
                row_cells = table.add_row().cells
                for i in range(len(column_names)):
                    row_cells[i].text = str(row[i])

            # saves the doc to a new file path
            doc.save(filepath + "/" + filename + ".docx")
            os.startfile(filepath + "/" + filename + ".docx")

    def save_doc_personnel(self, filepath, filename, date1, date2):
        # populate the report tree
        cont_inue = self.treeview.populate_personnel_report_bydate(date1, date2)
        if cont_inue:
            # open an existing Word document
            doc = docx.Document(".\Documents\Personnel Report Template.docx")

            column_names = [
                "Personnel Number",
                "First Name",
                "Last Name",
                "Personnel Type",
                "Date",
                "Time In",
                "Time Out",
            ]
            table = doc.add_table(rows=1, cols=len(column_names))
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells

            for i in range(len(column_names)):
                hdr_cells[i].text = column_names[i]

            treeview_data = []
            # get the table data of student
            for child in self.treeview.personnel_report_tree.get_children():
                values = self.treeview.personnel_report_tree.item(child)["values"]
                if values not in treeview_data:
                    treeview_data.append(values)

            # Inserts the table data of student
            for row in treeview_data:
                row_cells = table.add_row().cells
                for i in range(len(column_names)):
                    row_cells[i].text = str(row[i])

            # saves the doc to a new file path
            doc.save(filepath + "/" + filename + ".docx")

    def print_doc_visitor(self, filepath, filename, date1, date2):
        # populate the report tree
        cont_inue = self.treeview.populate_visitor_report_bydate(date1, date2)
        if cont_inue:
            # open an existing Word document
            doc = docx.Document(".\Documents\Visitor Report Template.docx")

            column_names = [
                "Visitor Number",
                "First Name",
                "Last Name",
                "Date",
                "Time In",
                "Time Out",
            ]

            table = doc.add_table(rows=1, cols=len(column_names))
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells

            for i in range(len(column_names)):
                hdr_cells[i].text = column_names[i]

            treeview_data = []
            # get the table data of student
            for child in self.treeview.visitor_report_tree.get_children():
                values = self.treeview.visitor_report_tree.item(child)["values"]
                if values not in treeview_data:
                    treeview_data.append(values)

            # Inserts the table data of student
            for row in treeview_data:
                row_cells = table.add_row().cells
                for i in range(len(column_names)):
                    row_cells[i].text = str(row[i])

            # saves the doc to a new file path
            doc.save(filepath + "/" + filename + ".docx")
            os.startfile(filepath + "/" + filename + ".docx")

    def save_doc_visitor(self, filepath, filename, date1, date2):
        # populate the report tree
        cont_inue = self.treeview.populate_visitor_report_bydate(date1, date2)
        if cont_inue:
            # open an existing Word document
            doc = docx.Document(".\Documents\Visitor Report Template.docx")

            column_names = [
                "Visitor Number",
                "First Name",
                "Last Name",
                "Date",
                "Time In",
                "Time Out",
            ]

            table = doc.add_table(rows=1, cols=len(column_names))
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells

            for i in range(len(column_names)):
                hdr_cells[i].text = column_names[i]

            treeview_data = []
            # get the table data of student
            for child in self.treeview.visitor_report_tree.get_children():
                values = self.treeview.visitor_report_tree.item(child)["values"]
                if values not in treeview_data:
                    treeview_data.append(values)

            # Inserts the table data of student
            for row in treeview_data:
                row_cells = table.add_row().cells
                for i in range(len(column_names)):
                    row_cells[i].text = str(row[i])

            # saves the doc to a new file path
            doc.save(filepath + "/" + filename + ".docx")

    def print_pdf_student(self, filepath, filename, date1, date2):
        # populate the report tree
        cont_inue = self.treeview.populate_student_report_bydate(date1, date2)
        if cont_inue:
            # open an existing Word document
            doc = docx.Document(".\Documents\Student Report Template.docx")

            column_names = [
                "Student_No",
                "Firstname",
                "Lastname",
                "Program",
                "Section",
                "Date",
                "Time In",
                "Time Out",
            ]

            table = doc.add_table(rows=1, cols=len(column_names))
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells

            for i in range(len(column_names)):
                hdr_cells[i].text = column_names[i]

            treeview_data = []
            # get the table data of student
            for child in self.treeview.student_report_tree.get_children():
                values = self.treeview.student_report_tree.item(child)["values"]
                if values not in treeview_data:
                    treeview_data.append(values)

            # Inserts the table data of student
            for row in treeview_data:
                row_cells = table.add_row().cells
                for i in range(len(column_names)):
                    row_cells[i].text = str(row[i])

            # saves the doc to a new file path
            doc.save(filepath + "/" + filename + ".docx")

            docx_file = filepath + "/" + filename + ".docx"
            pdf_file = filepath + "/" + filename + ".pdf"

            docx2pdf.convert(docx_file, pdf_file)
            os.remove(docx_file)
            os.startfile(filepath + "/" + filename + ".pdf")

    # JOCRIAUS-23/04/2023-------------------------------------------------------------------------------------
    # yung 12 na other functions gawin mong parepareho, eto yung gawin mong basis
    def save_pdf_student(self, filepath, filename, date1, date2):
        # populate the report tree
        cont_inue = self.treeview.populate_student_report_bydate(date1, date2)
        if cont_inue:
            # open an existing Word document
            doc = docx.Document(".\Documents\Student Report Template.docx")

            column_names = [
                "Student_No",
                "Firstname",
                "Lastname",
                "Program",
                "Section",
                "Date",
                "Time In",
                "Time Out",
            ]

            table = doc.add_table(rows=1, cols=len(column_names))
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells

            for i in range(len(column_names)):
                hdr_cells[i].text = column_names[i]

            treeview_data = []
            # get the table data of student
            for child in self.treeview.student_report_tree.get_children():
                values = self.treeview.student_report_tree.item(child)["values"]
                if values not in treeview_data:
                    treeview_data.append(values)

            # Inserts the table data of student
            for row in treeview_data:
                row_cells = table.add_row().cells
                for i in range(len(column_names)):
                    row_cells[i].text = str(row[i])

            # saves the doc to a new file path
            doc.save(filepath + "/" + filename + ".docx")

            docx_file = filepath + "/" + filename + ".docx"
            pdf_file = filepath + "/" + filename + ".pdf"

            docx2pdf.convert(docx_file, pdf_file)
            os.remove(docx_file)

    # JOCRIAUS-23/04/2023-------------------------------------------------------------------------------------

    def print_pdf_personnel(self, filepath, filename, date1, date2):
        # populate the report tree
        cont_inue = self.treeview.populate_personnel_report_bydate(date1, date2)
        if cont_inue:
            # open an existing Word document
            doc = docx.Document(".\Documents\Personnel Report Template.docx")

            column_names = [
                "Personnel Number",
                "First Name",
                "Last Name",
                "Personnel Type",
                "Date",
                "Time In",
                "Time Out",
            ]
            table = doc.add_table(rows=1, cols=len(column_names))
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells

            for i in range(len(column_names)):
                hdr_cells[i].text = column_names[i]

            treeview_data = []
            # get the table data of student
            for child in self.treeview.personnel_report_tree.get_children():
                values = self.treeview.personnel_report_tree.item(child)["values"]
                if values not in treeview_data:
                    treeview_data.append(values)

            # Inserts the table data of student
            for row in treeview_data:
                row_cells = table.add_row().cells
                for i in range(len(column_names)):
                    row_cells[i].text = str(row[i])

            # saves the doc to a new file path
            doc.save(filepath + "/" + filename + ".docx")

            docx_file = filepath + "/" + filename + ".docx"
            pdf_file = filepath + "/" + filename + ".pdf"

            docx2pdf.convert(docx_file, pdf_file)
            os.remove(docx_file)
            os.startfile(filepath + "/" + filename + ".pdf")

    def save_pdf_personnel(self, filepath, filename, date1, date2):
        # populate the report tree
        cont_inue = self.treeview.populate_personnel_report_bydate(date1, date2)
        if cont_inue:
            # open an existing Word document
            doc = docx.Document(".\Documents\Personnel Report Template.docx")

            column_names = [
                "Personnel Number",
                "First Name",
                "Last Name",
                "Personnel Type",
                "Date",
                "Time In",
                "Time Out",
            ]
            table = doc.add_table(rows=1, cols=len(column_names))
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells

            for i in range(len(column_names)):
                hdr_cells[i].text = column_names[i]

            treeview_data = []
            # get the table data of student
            for child in self.treeview.personnel_report_tree.get_children():
                values = self.treeview.personnel_report_tree.item(child)["values"]
                if values not in treeview_data:
                    treeview_data.append(values)

            # Inserts the table data of student
            for row in treeview_data:
                row_cells = table.add_row().cells
                for i in range(len(column_names)):
                    row_cells[i].text = str(row[i])

            # saves the doc to a new file path
            doc.save(filepath + "/" + filename + ".docx")

            docx_file = filepath + "/" + filename + ".docx"
            pdf_file = filepath + "/" + filename + ".pdf"

            docx2pdf.convert(docx_file, pdf_file)
            os.remove(docx_file)

    def print_pdf_visitor(self, filepath, filename, date1, date2):
        # populate the report tree
        cont_inue = self.treeview.populate_visitor_report_bydate(date1, date2)
        if cont_inue:
            # open an existing Word document
            doc = docx.Document(".\Documents\Visitor Report Template.docx")

            column_names = [
                "Visitor Number",
                "First Name",
                "Last Name",
                "Date",
                "Time In",
                "Time Out",
            ]

            table = doc.add_table(rows=1, cols=len(column_names))
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells

            for i in range(len(column_names)):
                hdr_cells[i].text = column_names[i]

            treeview_data = []
            # get the table data of student
            for child in self.treeview.visitor_report_tree.get_children():
                values = self.treeview.visitor_report_tree.item(child)["values"]
                if values not in treeview_data:
                    treeview_data.append(values)

            # Inserts the table data of student
            for row in treeview_data:
                row_cells = table.add_row().cells
                for i in range(len(column_names)):
                    row_cells[i].text = str(row[i])

            # saves the doc to a new file path
            doc.save(filepath + "/" + filename + ".docx")

            docx_file = filepath + "/" + filename + ".docx"
            pdf_file = filepath + "/" + filename + ".pdf"

            docx2pdf.convert(docx_file, pdf_file)
            os.remove(docx_file)
            os.startfile(filepath + "/" + filename + ".pdf")

    def save_pdf_visitor(self, filepath, filename, date1, date2):
        # populate the report tree
        cont_inue =self.treeview.populate_visitor_report_bydate(date1, date2)
        if cont_inue:
            # open an existing Word document
            doc = docx.Document(".\Documents\Visitor Report Template.docx")

            column_names = [
                "Visitor Number",
                "First Name",
                "Last Name",
                "Date",
                "Time In",
                "Time Out",
            ]

            table = doc.add_table(rows=1, cols=len(column_names))
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells

            for i in range(len(column_names)):
                hdr_cells[i].text = column_names[i]

            treeview_data = []
            # get the table data of student
            for child in self.treeview.visitor_report_tree.get_children():
                values = self.treeview.visitor_report_tree.item(child)["values"]
                if values not in treeview_data:
                    treeview_data.append(values)

            # Inserts the table data of student
            for row in treeview_data:
                row_cells = table.add_row().cells
                for i in range(len(column_names)):
                    row_cells[i].text = str(row[i])

            # saves the doc to a new file path
            doc.save(filepath + "/" + filename + ".docx")

            docx_file = filepath + "/" + filename + ".docx"
            pdf_file = filepath + "/" + filename + ".pdf"

            docx2pdf.convert(docx_file, pdf_file)
            os.remove(docx_file)
